from __future__ import annotations

import logging
import re
from dataclasses import dataclass

from django.db import transaction
from django.utils import timezone
from docx import Document as DocxDocument
from pypdf import PdfReader

from apps.documents.models import Document, DocumentText

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """
    Error controlado durante extracción de texto.
    """

    pass


@dataclass(frozen=True)
class ExtractedTextResult:
    content: str
    word_count: int
    character_count: int
    extraction_engine: str


class DocumentTextExtractor:
    """
    Extrae texto desde documentos PDF o DOCX.
    """

    PDF_MIME = "application/pdf"
    DOCX_MIME = (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    )

    @transaction.atomic
    def extract_and_save(self, document: Document) -> DocumentText:
        result = self.extract(document=document)

        document_text, _ = DocumentText.objects.update_or_create(
            document=document,
            defaults={
                "content": result.content,
                "word_count": result.word_count,
                "character_count": result.character_count,
                "extraction_engine": result.extraction_engine,
                "extracted_at": timezone.now(),
            },
        )

        logger.info(
            "Texto extraído correctamente. document_id=%s words=%s",
            document.id,
            result.word_count,
        )

        return document_text

    def extract(self, document: Document) -> ExtractedTextResult:
        if document.mime_type == self.PDF_MIME:
            raw_text = self._extract_pdf(document=document)
            engine = "pypdf"

        elif document.mime_type == self.DOCX_MIME:
            raw_text = self._extract_docx(document=document)
            engine = "python-docx"

        else:
            raise TextExtractionError(
                "Formato no soportado para extracción de texto."
            )

        content = self._normalize_text(raw_text)

        if len(content) < 100:
            raise TextExtractionError(
                "No se pudo extraer suficiente texto del documento."
            )

        words = re.findall(r"\b\w+\b", content, flags=re.UNICODE)

        return ExtractedTextResult(
            content=content,
            word_count=len(words),
            character_count=len(content),
            extraction_engine=engine,
        )

    def _extract_pdf(self, document: Document) -> str:
        try:
            with document.original_file.open("rb") as file_obj:
                reader = PdfReader(file_obj)
                pages_text: list[str] = []

                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    pages_text.append(page_text)

                return "\n".join(pages_text)

        except Exception as exc:
            logger.exception(
                "Error extrayendo PDF. document_id=%s",
                document.id,
            )
            raise TextExtractionError(
                "No se pudo extraer texto del PDF."
            ) from exc

    def _extract_docx(self, document: Document) -> str:
        try:
            with document.original_file.open("rb") as file_obj:
                docx = DocxDocument(file_obj)
                parts: list[str] = []

                for paragraph in docx.paragraphs:
                    if paragraph.text.strip():
                        parts.append(paragraph.text.strip())

                for table in docx.tables:
                    for row in table.rows:
                        row_text = " ".join(
                            cell.text.strip()
                            for cell in row.cells
                            if cell.text.strip()
                        )
                        if row_text:
                            parts.append(row_text)

                return "\n".join(parts)

        except Exception as exc:
            logger.exception(
                "Error extrayendo DOCX. document_id=%s",
                document.id,
            )
            raise TextExtractionError(
                "No se pudo extraer texto del DOCX."
            ) from exc

    def _normalize_text(self, text: str) -> str:
        text = text.replace("\x00", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()